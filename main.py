import os
import fitz  # PyMuPDF, used to convert PDF pages to images
import easyocr  # For OCR (Bengali + English)
from docx import Document  # To create Word documents
import pandas as pd  # For Excel output
from tqdm import tqdm  # Progress bar
import re  # Regular expressions to detect Arabic text

# ---------- CONFIG ----------
INPUT_PDF = r"book.pdf"  # Path to the input PDF
OUTPUT_DIR = r"output"  # Folder to store all outputs
PAGES_TO_CONVERT = 20  # Number of pages to process
# -----------------------------

# Create necessary folders for images, Word, and Excel outputs
os.makedirs(OUTPUT_DIR, exist_ok=True)
IMAGES_DIR = os.path.join(OUTPUT_DIR, "images")  # Folder for page images
DOCS_DIR = os.path.join(OUTPUT_DIR, "word")  # Folder for Word outputs
EXCEL_DIR = os.path.join(OUTPUT_DIR, "excel")  # Folder for Excel outputs
os.makedirs(IMAGES_DIR, exist_ok=True)
os.makedirs(DOCS_DIR, exist_ok=True)
os.makedirs(EXCEL_DIR, exist_ok=True)

# 1) Convert PDF pages to images using PyMuPDF
print("Converting PDF pages to images...")
doc = fitz.open(INPUT_PDF)  # Open the PDF file
for i in range(min(PAGES_TO_CONVERT, len(doc))):
    page = doc[i]  # Select page
    pix = page.get_pixmap(dpi=300)  # Render page as high-resolution image
    img_path = os.path.join(IMAGES_DIR, f"page_{i+1:02d}.png")  # Image filename
    pix.save(img_path)  # Save image

# 2) OCR with EasyOCR (English + Bengali)
print("Running OCR (English + Bengali)...")
reader = easyocr.Reader(['bn', 'en'])  # Initialize OCR reader
page_texts = []  # Store extracted text per page

# Arabic Unicode regex to skip any Arabic text
arabic_re = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]')

def clean_text(text_lines):
    """
    Merge short lines into paragraphs, remove Arabic text and empty lines.
    Returns a list of cleaned paragraphs.
    """
    clean_lines = []
    buffer = ""  # Temporary buffer to merge short lines
    for line in text_lines:
        line = line.strip()
        if not line or arabic_re.search(line):
            continue  # Skip empty lines or Arabic text
        if len(line) < 50:
            buffer += line + " "  # Merge short lines
        else:
            if buffer:
                clean_lines.append(buffer.strip())  # Add merged paragraph
                buffer = ""
            clean_lines.append(line)
    if buffer:
        clean_lines.append(buffer.strip())
    return clean_lines

def split_columns(paragraphs):
    """
    Split paragraphs into Title/Description columns if ':' or '-' is present.
    Returns a list of dictionaries for Excel rows.
    """
    rows = []
    for para in paragraphs:
        if ':' in para:
            parts = para.split(':', 1)
            rows.append({'Title': parts[0].strip(), 'Description': parts[1].strip(), 'FullText': para})
        elif '-' in para:
            parts = para.split('-', 1)
            rows.append({'Title': parts[0].strip(), 'Description': parts[1].strip(), 'FullText': para})
        else:
            rows.append({'Title': para, 'Description': '', 'FullText': para})
    return rows

# OCR each page
for i in tqdm(range(1, PAGES_TO_CONVERT + 1), desc="Processing pages"):
    img_path = os.path.join(IMAGES_DIR, f"page_{i:02d}.png")  # Image path
    results = reader.readtext(img_path, detail=0)  # OCR text extraction
    paragraphs = clean_text(results)  # Clean extracted text
    page_texts.append(paragraphs)  # Add page paragraphs to list

# 3) Save to Word
print("Creating Word file...")
docx_file = Document()  # Create new Word document
for i, paragraphs in enumerate(page_texts, start=1):
    heading = docx_file.add_heading(f"Page {i}", level=2)  # Add page heading
    heading.runs[0].bold = True  # Make heading bold
    for para in paragraphs:
        docx_file.add_paragraph(para)  # Add each paragraph
    docx_file.add_page_break()  # Page break after each page

word_out = os.path.join(DOCS_DIR, "Output.docx")  # Word output path
docx_file.save(word_out)
print(f"Word file saved: {word_out}")

# 4) Save to Excel
print("Saving text to Excel...")
all_rows = []
for i, paragraphs in enumerate(page_texts, start=1):
    rows = split_columns(paragraphs)  # Split into columns for Excel
    for row in rows:
        row['PageNumber'] = i  # Add page number
        all_rows.append(row)

df = pd.DataFrame(all_rows, columns=['PageNumber', 'Title', 'Description', 'FullText'])  # Create DataFrame
excel_out = os.path.join(EXCEL_DIR, "Output.xlsx")  # Excel output path
df.to_excel(excel_out, index=False)  # Save Excel file
print(f"Excel file saved: {excel_out}")

print("OCR extraction completed successfully.")  # Final message
